import io

import fitz
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.parsers.resume_parser import ResumeLink, ResumeParser


def add_hyperlink(document: Document, label: str, url: str) -> None:
    paragraph = document.add_paragraph()
    relationship_id = paragraph.part.relate_to(
        url,
        RT.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def test_extracts_hidden_docx_hyperlink_destination() -> None:
    document = Document()
    document.add_paragraph("Rahul Sharma")
    add_hyperlink(
        document,
        "rahul51",
        "https://www.linkedin.com/in/rahul51/?utm_source=resume",
    )
    output = io.BytesIO()
    document.save(output)

    text, _ = ResumeParser.parse_docx(output.getvalue())

    assert "LinkedIn: https://www.linkedin.com/in/rahul51" in text
    assert "utm_source" not in text


def test_extracts_link_from_docx_header() -> None:
    document = Document()
    document.add_paragraph("Rahul Sharma")
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    relationship_id = paragraph.part.relate_to(
        "https://github.com/rahul51",
        RT.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    value = OxmlElement("w:t")
    value.text = "rahul51"
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    output = io.BytesIO()
    document.save(output)

    text, _ = ResumeParser.parse_docx(output.getvalue())

    assert "GitHub: https://github.com/rahul51" in text


def test_extracts_hidden_pdf_hyperlink_destination() -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Rahul Sharma")
    page.insert_text((72, 96), "rahul51")
    page.insert_link(
        {
            "kind": fitz.LINK_URI,
            "from": fitz.Rect(70, 82, 150, 102),
            "uri": "https://github.com/rahul51",
        }
    )
    content = document.tobytes()
    document.close()

    text, _ = ResumeParser.parse_pdf(content)

    assert "GitHub: https://github.com/rahul51" in text


def test_rejects_unsafe_and_private_link_destinations() -> None:
    text = ResumeParser._merge_links(
        "Rahul Sharma",
        [
            ResumeLink("profile", "javascript:alert(1)"),
            ResumeLink("internal", "http://127.0.0.1/admin"),
        ],
    )

    assert text == "Rahul Sharma"

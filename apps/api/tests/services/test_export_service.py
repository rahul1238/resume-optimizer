import io

import fitz
import pytest
from docx import Document

from app.ai.schemas import BulletRewrite, ResumeImprovementResult
from app.services.export_service import (
    ExportContext,
    ResumeExportService,
    ResumePageLimitError,
)
from app.services.resume_storage_service import ResumeStorageService

SAMPLE_DRAFT = """Jordan Lee
jordan@example.com | +1 555 0100 | Austin, TX | https://linkedin.com/in/jordanlee

PROFESSIONAL SUMMARY
Backend engineer building reliable Python services and APIs.

EXPERIENCE
Senior Software Engineer | Example Tech | 2023-Present
- Improved API availability to 99.95% through observability and testing.
- Reduced deployment time by 40% using automated delivery pipelines.

SKILLS
Python, FastAPI, PostgreSQL, Docker, Google Cloud

EDUCATION
B.S. Computer Science | State University
"""


def improvement_result() -> ResumeImprovementResult:
    return ResumeImprovementResult(
        optimized_resume_draft=SAMPLE_DRAFT,
        suggested_summary="Backend engineer building reliable Python services.",
        summary_reason="More concise.",
        bullet_rewrites=[
            BulletRewrite(
                original="Built APIs",
                suggested="Built reliable APIs serving 10,000 daily requests",
                reason="Adds scope.",
            )
        ],
        ats_recommendations=["Use standard section headings."],
    )


def test_structure_preserves_semantic_hierarchy() -> None:
    resume = ResumeExportService.structure(SAMPLE_DRAFT)

    assert resume.header[0] == "Jordan Lee"
    assert [section.heading for section in resume.sections] == [
        "PROFESSIONAL SUMMARY",
        "EXPERIENCE",
        "SKILLS",
        "EDUCATION",
    ]
    assert ResumeExportService.is_bullet(resume.sections[1].lines[1])


def test_ats_exports_are_text_parseable_and_fit_one_page() -> None:
    pdf = ResumeExportService.to_pdf(SAMPLE_DRAFT, target_pages=1)
    docx = ResumeExportService.to_docx(SAMPLE_DRAFT, target_pages=1)

    with fitz.open(stream=pdf, filetype="pdf") as document:
        assert len(document) == 1
        assert "99.95%" in "\n".join(page.get_text() for page in document)
        assert document[0].get_links()[0]["uri"] == (
            "https://linkedin.com/in/jordanlee"
        )
    word_document = Document(io.BytesIO(docx))
    text = "\n".join(paragraph.text for paragraph in word_document.paragraphs)
    assert "PROFESSIONAL SUMMARY" in text
    assert "99.95%" in text
    assert any(
        relationship.target_ref == "https://linkedin.com/in/jordanlee"
        for relationship in word_document.part.rels.values()
        if relationship.is_external
    )


def test_preserve_docx_keeps_style_and_applies_exact_rewrite(
    monkeypatch,
) -> None:
    original = Document()
    paragraph = original.add_paragraph()
    paragraph.style = original.styles["Normal"]
    run = paragraph.add_run("Built APIs")
    run.bold = True
    source = io.BytesIO()
    original.save(source)
    monkeypatch.setattr(
        ResumeStorageService,
        "read_bytes",
        lambda _path: source.getvalue(),
    )
    context = ExportContext(
        improvement_result(),
        "docx",
        "original.docx",
        SAMPLE_DRAFT,
        source.getvalue(),
    )

    exported = Document(io.BytesIO(ResumeExportService.preserve_docx(context)))

    assert exported.paragraphs[0].text == (
        "Built reliable APIs serving 10,000 daily requests"
    )
    assert exported.paragraphs[0].runs[0].bold is True


def test_restores_links_missing_from_cached_optimized_draft() -> None:
    cached_draft = SAMPLE_DRAFT.replace(
        " | https://linkedin.com/in/jordanlee",
        "",
    )
    source_text = """Jordan Lee
rahul51

LINKS
LinkedIn: https://www.linkedin.com/in/rahul51
GitHub: https://github.com/rahul51
"""

    restored = ResumeExportService._restore_source_links(cached_draft, source_text)
    resume = ResumeExportService.structure(restored)

    assert any("linkedin.com/in/rahul51" in line for line in resume.header)
    assert any("github.com/rahul51" in line for line in resume.header)


def test_one_page_target_never_silently_returns_two_pages(
    monkeypatch,
) -> None:
    oversized = fitz.open()
    oversized.new_page()
    oversized.new_page()
    content = oversized.tobytes()
    oversized.close()
    monkeypatch.setattr(
        ResumeExportService,
        "_render_pdf",
        lambda *_args: content,
    )

    with pytest.raises(ResumePageLimitError):
        ResumeExportService.to_pdf(SAMPLE_DRAFT, target_pages=1)

import io

import fitz
import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.auth.dependencies import CurrentUser, get_current_user
from app.main import app
from app.repositories.resume_repository import ResumeRepository
from app.services.resume_storage_service import ResumeStorageService, StoredResume

client = TestClient(app)


@pytest.fixture(autouse=True)
def override_authentication() -> None:
    async def get_test_user() -> CurrentUser:
        return CurrentUser(
            uid="test-user-id",
            email="test@example.com",
            email_verified=True,
        )

    app.dependency_overrides[get_current_user] = get_test_user
    yield
    app.dependency_overrides.clear()


@pytest.fixture(autouse=True)
def override_storage(monkeypatch: pytest.MonkeyPatch) -> None:
    def store_original(**_kwargs: object) -> StoredResume:
        return StoredResume(
            resume_id="b6b6b661-144a-4cb7-b0d8-2ce2912e211d",
            storage_path=(
                "resumes/test-user-id/b6b6b661-144a-4cb7-b0d8-2ce2912e211d/original.pdf"
            ),
            text_storage_path=(
                "resumes/test-user-id/b6b6b661-144a-4cb7-b0d8-2ce2912e211d/"
                "extracted.txt"
            ),
        )

    monkeypatch.setattr(ResumeStorageService, "store_original", store_original)
    monkeypatch.setattr(ResumeRepository, "create", lambda _record: None)
    monkeypatch.setattr(
        ResumeRepository,
        "find_owned_by_hash",
        lambda _owner_uid, _content_sha256: None,
    )


def test_upload_pdf_extracts_text() -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Ada Lovelace\nPython Engineer")
    content = document.tobytes()
    document.close()

    response = client.post(
        "/api/v1/resumes/upload",
        files={"file": ("resume.pdf", content, "application/pdf")},
    )

    assert response.status_code == 200
    assert response.json()["file_type"] == "pdf"
    assert response.json()["page_count"] == 1
    assert response.json()["resume_id"] == "b6b6b661-144a-4cb7-b0d8-2ce2912e211d"
    assert response.json()["storage_path"].startswith("resumes/test-user-id/")
    assert "Ada Lovelace" in response.json()["text"]


def test_upload_docx_extracts_text() -> None:
    document = Document()
    document.add_paragraph("Grace Hopper")
    document.add_paragraph("Compiler Pioneer")
    content = io.BytesIO()
    document.save(content)

    response = client.post(
        "/api/v1/resumes/upload",
        files={
            "file": (
                "resume.docx",
                content.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    assert response.json()["file_type"] == "docx"
    assert response.json()["page_count"] is None
    assert "Grace Hopper" in response.json()["text"]


def test_upload_rejects_unsupported_format() -> None:
    response = client.post(
        "/api/v1/resumes/upload",
        files={"file": ("resume.txt", b"Resume", "text/plain")},
    )

    assert response.status_code == 415
    assert response.json() == {
        "detail": "Only PDF and DOCX resumes are supported.",
        "code": "unsupported_resume_format",
    }


def test_upload_rejects_a_file_with_a_spoofed_pdf_extension() -> None:
    response = client.post(
        "/api/v1/resumes/upload",
        files={"file": ("resume.pdf", b"not a pdf", "application/pdf")},
    )

    assert response.status_code == 422
    assert response.json()["code"] == "unreadable_resume"


def test_upload_rejects_a_file_larger_than_the_allowed_limit() -> None:
    content = b"0" * (5 * 1024 * 1024 + 1)

    response = client.post(
        "/api/v1/resumes/upload",
        files={"file": ("resume.pdf", content, "application/pdf")},
    )

    assert response.status_code == 413
    assert response.json()["code"] == "resume_too_large"


def test_upload_requires_a_firebase_id_token() -> None:
    app.dependency_overrides.pop(get_current_user)

    response = client.post(
        "/api/v1/resumes/upload",
        files={"file": ("resume.pdf", b"%PDF-", "application/pdf")},
    )

    assert response.status_code == 401
    assert response.json() == {
        "detail": "A Firebase ID token is required.",
        "code": "missing_authentication",
    }


def test_upload_cors_preflight_allows_the_local_frontend() -> None:
    response = client.options(
        "/api/v1/resumes/upload",
        headers={
            "Origin": "http://localhost:3000",
            "Access-Control-Request-Method": "POST",
            "Access-Control-Request-Headers": "authorization,content-type",
        },
    )

    assert response.status_code == 200
    assert response.headers["access-control-allow-origin"] == "http://localhost:3000"
    assert "POST" in response.headers["access-control-allow-methods"]

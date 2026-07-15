import io
import logging
import re
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal

import fitz
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.ai.schemas import ResumeImprovementResult
from app.core.config import settings
from app.parsers.resume_parser import ResumeParser, ResumeParsingError
from app.repositories.resume_repository import ResumeRepository
from app.services.improvement_service import ImprovementService
from app.services.resume_storage_service import ResumeStorageService

ExportMode = Literal["ats", "preserve"]
logger = logging.getLogger(__name__)


class ResumeExportError(Exception):
    status_code = 409
    code = "resume_draft_not_ready"
    message = "Generate and save an optimized resume draft before exporting."

    def __init__(self, message: str | None = None) -> None:
        self.message = message or self.message
        super().__init__(self.message)


class ResumeExportValidationError(ResumeExportError):
    code = "resume_export_validation_failed"
    message = "The generated resume could not be validated for reliable text parsing."


class ResumePageLimitError(ResumeExportError):
    code = "resume_page_limit_exceeded"
    message = (
        "The complete resume cannot fit within the selected page limit without "
        "hurting readability. Shorten the draft or select two pages."
    )


class LatexCompilationError(ResumeExportError):
    status_code = 503
    code = "latex_compiler_unavailable"
    message = "The PDF generator is temporarily unavailable."


@dataclass(frozen=True)
class ResumeSection:
    heading: str
    lines: list[str] = field(default_factory=list)


@dataclass(frozen=True)
class StructuredResume:
    header: list[str]
    sections: list[ResumeSection]

    @property
    def all_text(self) -> str:
        parts = list(self.header)
        for section in self.sections:
            parts.append(section.heading)
            parts.extend(section.lines)
        return "\n".join(parts)


@dataclass(frozen=True)
class LayoutProfile:
    body_size: float
    leading: float
    heading_size: float
    name_size: float
    margin: float
    section_space: float
    line_space: float


@dataclass(frozen=True)
class ExportContext:
    result: ResumeImprovementResult
    file_type: str
    original_storage_path: str
    draft: str
    original_content: bytes


class ResumeExportService:
    section_names = {
        "summary",
        "professional summary",
        "profile",
        "objective",
        "experience",
        "professional experience",
        "work experience",
        "employment history",
        "education",
        "skills",
        "technical skills",
        "core competencies",
        "projects",
        "certifications",
        "licenses",
        "awards",
        "publications",
        "volunteering",
        "languages",
    }
    profiles = (
        LayoutProfile(10, 13, 11, 17, 0.68, 7, 2.5),
        LayoutProfile(9, 12, 10.5, 16, 0.58, 5, 1.5),
        LayoutProfile(9, 11, 10, 15, 0.5, 4, 1),
    )

    @classmethod
    def get_context(cls, owner_uid: str, analysis_id: str) -> ExportContext:
        improvement = ImprovementService.get(owner_uid, analysis_id)
        result = ImprovementService.result(improvement)
        if not result.optimized_resume_draft.strip():
            raise ResumeExportError()
        resume = ResumeRepository.get_owned(improvement.resume_id, owner_uid)
        original_content = ResumeStorageService.read_bytes(
            resume.original_storage_path
        )
        try:
            if resume.file_type == "pdf":
                source_text, _ = ResumeParser.parse_pdf(original_content)
            else:
                source_text, _ = ResumeParser.parse_docx(original_content)
        except ResumeParsingError:
            logger.warning(
                "Original resume could not be reparsed during export",
                extra={"resume_id": resume.resume_id},
            )
            source_text = ResumeStorageService.read_text(resume.text_storage_path)
        draft = cls._restore_source_links(
            result.optimized_resume_draft,
            source_text,
        )
        return ExportContext(
            result,
            resume.file_type,
            resume.original_storage_path,
            draft,
            original_content,
        )

    @classmethod
    def get_draft(cls, owner_uid: str, analysis_id: str) -> str:
        return cls.get_context(owner_uid, analysis_id).draft

    @classmethod
    def _restore_source_links(cls, draft: str, source_text: str) -> str:
        pattern = re.compile(r"(?:https?://|mailto:|tel:)[^\s<>|]+")
        draft_urls = {
            normalized
            for value in pattern.findall(draft)
            if (normalized := ResumeParser._normalize_url(value.rstrip(".,;)")))
        }
        missing: list[tuple[str, str]] = []
        seen = set(draft_urls)
        for line in source_text.splitlines():
            for match in pattern.finditer(line):
                normalized = ResumeParser._normalize_url(
                    match.group(0).rstrip(".,;)")
                )
                if not normalized or normalized in seen:
                    continue
                seen.add(normalized)
                prefix = line[: match.start()].strip().rstrip(":|- ")
                label = ResumeParser._link_label(prefix, normalized)
                missing.append((label, normalized))
        if not missing:
            return draft.strip()

        link_line = " | ".join(f"{label}: {url}" for label, url in missing)
        lines = draft.strip().splitlines()
        heading_index = next(
            (index for index, line in enumerate(lines) if cls.is_heading(line)),
            len(lines),
        )
        lines.insert(heading_index, link_line)
        return "\n".join(lines)

    @classmethod
    def is_heading(cls, line: str) -> bool:
        normalized = line.strip().rstrip(":").lower()
        return normalized in cls.section_names or (
            len(line) <= 45
            and line.isupper()
            and any(character.isalpha() for character in line)
        )

    @staticmethod
    def is_bullet(line: str) -> bool:
        return bool(re.match(r"^\s*[-*\u2022]\s+", line))

    @staticmethod
    def bullet_text(line: str) -> str:
        return re.sub(r"^\s*[-*\u2022]\s+", "", line).strip()

    @classmethod
    def structure(cls, draft: str) -> StructuredResume:
        header: list[str] = []
        sections: list[ResumeSection] = []
        heading: str | None = None
        lines: list[str] = []

        def flush() -> None:
            nonlocal lines
            if heading is not None:
                sections.append(ResumeSection(heading.rstrip(":"), lines))
            lines = []

        for raw_line in draft.splitlines():
            line = re.sub(r"\s+", " ", raw_line).strip()
            if not line:
                continue
            if cls.is_heading(line):
                flush()
                heading = line
            elif heading is None:
                header.append(line)
            else:
                lines.append(line)
        flush()

        if not header and sections:
            first = sections.pop(0)
            header = [first.heading, *first.lines[:1]]
            if len(first.lines) > 1:
                sections.insert(0, ResumeSection("Profile", first.lines[1:]))
        if not header:
            raise ResumeExportError("The optimized draft does not contain usable text.")
        return StructuredResume(header[:6], sections)

    @classmethod
    def to_docx(
        cls,
        draft: str,
        target_pages: int = 1,
        profile: LayoutProfile | None = None,
    ) -> bytes:
        resume = cls.structure(draft)
        chosen = profile or cls._profile_for(resume, target_pages)
        document = Document()
        section = document.sections[0]
        section.section_start = WD_SECTION.CONTINUOUS
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(chosen.margin)
        section.bottom_margin = Inches(chosen.margin)
        section.left_margin = Inches(chosen.margin + 0.08)
        section.right_margin = Inches(chosen.margin + 0.08)

        normal = document.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(chosen.body_size)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        normal.paragraph_format.space_after = Pt(chosen.line_space)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        name = document.add_paragraph()
        name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        name.paragraph_format.space_after = Pt(2)
        run = name.add_run(resume.header[0])
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(chosen.name_size)
        for line in resume.header[1:]:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.paragraph_format.space_after = Pt(1)
            cls._add_docx_text(paragraph, line)

        for item in resume.sections:
            heading = document.add_paragraph()
            heading.paragraph_format.space_before = Pt(chosen.section_space)
            heading.paragraph_format.space_after = Pt(2)
            heading.paragraph_format.keep_with_next = True
            heading_run = heading.add_run(item.heading.upper())
            heading_run.bold = True
            heading_run.font.name = "Arial"
            heading_run.font.size = Pt(chosen.heading_size)
            heading_run.font.color.rgb = RGBColor(25, 25, 25)
            cls._add_bottom_border(heading)

            for line in item.lines:
                if cls.is_bullet(line):
                    paragraph = document.add_paragraph()
                    paragraph.style = document.styles["Normal"]
                    paragraph.paragraph_format.left_indent = Inches(0.16)
                    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
                    paragraph.add_run("\u2022 ")
                    cls._add_docx_text(paragraph, cls.bullet_text(line))
                else:
                    paragraph = document.add_paragraph()
                    cls._add_docx_text(paragraph, line)
                    paragraph.paragraph_format.keep_together = True

        output = io.BytesIO()
        document.save(output)
        content = output.getvalue()
        cls._validate_docx(content, resume.all_text)
        return content

    @classmethod
    def to_pdf(cls, draft: str, target_pages: int = 1) -> bytes:
        resume = cls.structure(draft)
        content, _ = cls._fit_pdf(resume, target_pages)
        cls._validate_pdf(content, resume.all_text)
        return content

    @classmethod
    def to_pdf_preview(cls, draft: str, target_pages: int = 1) -> tuple[bytes, int]:
        resume = cls.structure(draft)
        content, _, page_count = cls._best_fit_pdf(resume, target_pages)
        cls._validate_pdf(content, resume.all_text)
        return content, page_count

    @classmethod
    def preserve_docx(cls, context: ExportContext) -> bytes:
        if context.file_type != "docx":
            return cls.to_docx(context.result.optimized_resume_draft)
        document = Document(io.BytesIO(context.original_content))
        replacements = {
            rewrite.original.strip(): rewrite.suggested.strip()
            for rewrite in context.result.bullet_rewrites
            if rewrite.original.strip() and rewrite.suggested.strip()
        }
        for paragraph in cls._all_docx_paragraphs(document):
            original_text = paragraph.text.strip()
            replacement = replacements.get(original_text)
            if replacement:
                cls._replace_paragraph_text(paragraph, replacement)
        output = io.BytesIO()
        document.save(output)
        content = output.getvalue()
        cls._validate_docx(content, "\n".join(replacements.values()))
        return content

    @classmethod
    def _profile_for(cls, resume: StructuredResume, target_pages: int) -> LayoutProfile:
        _, profile = cls._fit_pdf(resume, target_pages)
        return profile

    @classmethod
    def _fit_pdf(
        cls, resume: StructuredResume, target_pages: int
    ) -> tuple[bytes, LayoutProfile]:
        content, profile, page_count = cls._best_fit_pdf(resume, target_pages)
        if page_count > target_pages:
            raise ResumePageLimitError()
        return content, profile

    @classmethod
    def _best_fit_pdf(
        cls, resume: StructuredResume, target_pages: int
    ) -> tuple[bytes, LayoutProfile, int]:
        last_content = b""
        last_page_count = 0
        for profile in cls.profiles:
            last_content = cls._render_pdf(resume, profile)
            with fitz.open(stream=last_content, filetype="pdf") as document:
                last_page_count = len(document)
                if last_page_count <= target_pages:
                    return last_content, profile, last_page_count
        return last_content, cls.profiles[-1], last_page_count

    @classmethod
    def _render_pdf(cls, resume: StructuredResume, profile: LayoutProfile) -> bytes:
        source = cls._latex_source(resume, profile)
        try:
            with tempfile.TemporaryDirectory(prefix="resume-latex-") as directory:
                source_path = Path(directory) / "resume.tex"
                source_path.write_text(source, encoding="utf-8")
                command = [settings.tectonic_binary]
                if settings.tectonic_only_cached:
                    command.append("--only-cached")
                command.extend(
                    ["--untrusted", "--outdir", directory, str(source_path)]
                )
                completed = subprocess.run(
                    command,
                    capture_output=True,
                    check=False,
                    text=True,
                    timeout=settings.latex_compile_timeout_seconds,
                )
                output_path = Path(directory) / "resume.pdf"
                if completed.returncode != 0 or not output_path.exists():
                    logger.error(
                        "Tectonic resume compilation failed: %s",
                        completed.stderr[-2000:],
                    )
                    raise LatexCompilationError()
                return output_path.read_bytes()
        except (OSError, subprocess.SubprocessError) as error:
            raise LatexCompilationError() from error

    @classmethod
    def _latex_source(
        cls,
        resume: StructuredResume,
        profile: LayoutProfile,
    ) -> str:
        header_lines = "\\\\\n".join(
            cls._latex_text(line) for line in resume.header[1:]
        )
        sections = "\n".join(cls._latex_section(section) for section in resume.sections)
        section_before = profile.section_space / 2.835
        heading_style = (
            rf"\bfseries\fontsize{{{profile.heading_size}}}"
            rf"{{{profile.heading_size + 2}}}\selectfont"
        )
        name_style = (
            rf"\bfseries\fontsize{{{profile.name_size}}}"
            rf"{{{profile.name_size + 2}}}\selectfont"
        )
        math_sizes = "\n".join(
            rf"\DeclareMathSizes{{{size}}}{{10}}{{7}}{{5}}"
            for size in {
                profile.body_size,
                profile.heading_size,
                profile.name_size,
            }
        )
        name = cls._latex_text(resume.header[0])
        return rf"""\documentclass[10pt,letterpaper]{{article}}
\usepackage[margin={profile.margin}in]{{geometry}}
\usepackage{{fontspec}}
\usepackage{{titlesec}}
\usepackage{{enumitem}}
\usepackage[hidelinks]{{hyperref}}
\usepackage{{xurl}}
\setmainfont{{lmsans10-regular.otf}}[BoldFont=lmsans10-bold.otf]
\setmonofont{{lmmono10-regular.otf}}
{math_sizes}
\pagestyle{{empty}}
\setlength{{\parindent}}{{0pt}}
\setlength{{\parskip}}{{{profile.line_space}pt}}
\setlist[itemize]{{leftmargin=1.15em,itemsep={profile.line_space}pt,topsep=1pt,parsep=0pt}}
\titleformat{{\section}}{{{heading_style}}}{{}}{{0pt}}{{\MakeUppercase}}[\vspace{{-2pt}}\titlerule]
\titlespacing*{{\section}}{{0pt}}{{{section_before:.2f}mm}}{{1.2mm}}
\begin{{document}}
\fontsize{{{profile.body_size}}}{{{profile.leading}}}\selectfont
\begin{{center}}
{{{name_style} {name}}}\\[1pt]
{header_lines}
\end{{center}}
\vspace{{-5pt}}
{sections}
\end{{document}}
"""

    @classmethod
    def _latex_section(cls, section: ResumeSection) -> str:
        output = [rf"\section*{{{cls._latex_escape(section.heading)}}}"]
        in_list = False
        for index, line in enumerate(section.lines):
            if cls.is_bullet(line):
                if not in_list:
                    output.append(r"\begin{itemize}")
                    in_list = True
                output.append(rf"\item {cls._latex_text(cls.bullet_text(line))}")
                continue
            if in_list:
                output.append(r"\end{itemize}")
                in_list = False
            next_is_bullet = (
                index + 1 < len(section.lines)
                and cls.is_bullet(section.lines[index + 1])
            )
            value = cls._latex_text(line)
            if next_is_bullet:
                output.append(rf"\textbf{{{value}}}\par")
            else:
                output.append(rf"{value}\par")
        if in_list:
            output.append(r"\end{itemize}")
        return "\n".join(output)

    @staticmethod
    def _add_docx_text(paragraph, text: str) -> None:
        link_pattern = re.compile(r"(?:https?://|mailto:|tel:)[^\s<>]+")
        position = 0
        for match in link_pattern.finditer(text):
            if match.start() > position:
                paragraph.add_run(text[position : match.start()])
            url = match.group(0)
            relationship_id = paragraph.part.relate_to(
                url,
                RT.HYPERLINK,
                is_external=True,
            )
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), relationship_id)
            run = OxmlElement("w:r")
            properties = OxmlElement("w:rPr")
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "0563C1")
            underline = OxmlElement("w:u")
            underline.set(qn("w:val"), "single")
            properties.extend((color, underline))
            run.append(properties)
            value = OxmlElement("w:t")
            value.text = url
            run.append(value)
            hyperlink.append(run)
            paragraph._p.append(hyperlink)
            position = match.end()
        if position < len(text):
            paragraph.add_run(text[position:])

    @classmethod
    def _latex_text(cls, text: str) -> str:
        link_pattern = re.compile(r"(?:https?://|mailto:|tel:)[^\s<>]+")
        parts: list[str] = []
        position = 0
        for match in link_pattern.finditer(text):
            parts.append(cls._latex_escape(text[position : match.start()]))
            url = match.group(0).replace("{", "%7B").replace("}", "%7D")
            parts.append(rf"\href{{\detokenize{{{url}}}}}{{\nolinkurl{{{url}}}}}")
            position = match.end()
        parts.append(cls._latex_escape(text[position:]))
        return "".join(parts)

    @staticmethod
    def _latex_escape(text: str) -> str:
        replacements = {
            "\\": r"\textbackslash{}",
            "&": r"\&",
            "%": r"\%",
            "$": r"\$",
            "#": r"\#",
            "_": r"\_",
            "{": r"\{",
            "}": r"\}",
            "~": r"\textasciitilde{}",
            "^": r"\textasciicircum{}",
        }
        return "".join(replacements.get(character, character) for character in text)

    @staticmethod
    def _add_bottom_border(paragraph) -> None:
        properties = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "666666")
        borders.append(bottom)
        properties.append(borders)

    @staticmethod
    def _all_docx_paragraphs(document):
        yield from document.paragraphs
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    @staticmethod
    def _replace_paragraph_text(paragraph, text: str) -> None:
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(text)

    @classmethod
    def _validate_docx(cls, content: bytes, expected: str) -> None:
        document = Document(io.BytesIO(content))
        extracted = "\n".join(
            paragraph.text for paragraph in cls._all_docx_paragraphs(document)
        )
        cls._validate_text(expected, extracted)

    @classmethod
    def _validate_pdf(cls, content: bytes, expected: str) -> None:
        with fitz.open(stream=content, filetype="pdf") as document:
            extracted = "\n".join(page.get_text("text") for page in document)
        cls._validate_text(expected, extracted)

    @staticmethod
    def _validate_text(expected: str, extracted: str) -> None:
        expected_words = set(re.findall(r"[a-z0-9+#.]{2,}", expected.lower()))
        extracted_words = set(re.findall(r"[a-z0-9+#.]{2,}", extracted.lower()))
        if not expected_words:
            return
        coverage = len(expected_words & extracted_words) / len(expected_words)
        if coverage < 0.95:
            raise ResumeExportValidationError()

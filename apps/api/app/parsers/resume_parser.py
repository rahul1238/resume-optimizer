import io
import ipaddress
import re
import zipfile
from dataclasses import dataclass
from urllib.parse import parse_qsl, urlencode, urlsplit, urlunsplit

import fitz
from docx import Document
from docx.oxml.ns import qn

from app.core.config import settings


class ResumeParsingError(Exception):
    """Raised when a resume file cannot be parsed."""


@dataclass(frozen=True)
class ResumeLink:
    label: str
    url: str


class ResumeParser:
    @staticmethod
    def parse_pdf(content: bytes) -> tuple[str, int]:
        if not content.startswith(b"%PDF-"):
            raise ResumeParsingError("The file contents are not a valid PDF.")

        try:
            document = fitz.open(stream=content, filetype="pdf")
        except (fitz.FileDataError, RuntimeError, ValueError) as error:
            raise ResumeParsingError("The PDF file could not be read.") from error

        try:
            if document.is_encrypted:
                raise ResumeParsingError("Password-protected PDFs are not supported.")

            pages = [page.get_text("text").strip() for page in document]
            links = [
                ResumeLink(
                    label=ResumeParser._pdf_link_label(page, link),
                    url=link.get("uri", ""),
                )
                for page in document
                for link in page.get_links()
                if link.get("kind") == fitz.LINK_URI
            ]
            text = ResumeParser._merge_links("\n".join(pages), links)
            return text, len(document)
        finally:
            document.close()

    @staticmethod
    def parse_docx(content: bytes) -> tuple[str, None]:
        ResumeParser._validate_docx_archive(content)

        try:
            document = Document(io.BytesIO(content))
        except Exception as error:
            raise ResumeParsingError("The DOCX file could not be read.") from error

        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        header_footer_text: list[str] = []
        for story in ResumeParser._docx_header_footers(document):
            header_footer_text.extend(paragraph.text for paragraph in story.paragraphs)
            header_footer_text.extend(
                cell.text
                for table in story.tables
                for row in table.rows
                for cell in row.cells
            )
        text = "\n".join(paragraphs + table_cells + header_footer_text).strip()
        links = ResumeParser._docx_links(document)
        text = ResumeParser._merge_links(text, links)

        # DOCX files do not store a reliable rendered page count.
        return text, None

    @staticmethod
    def _validate_docx_archive(content: bytes) -> None:
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                names = set(archive.namelist())
                uncompressed_size = sum(item.file_size for item in archive.infolist())
        except (zipfile.BadZipFile, OSError) as error:
            raise ResumeParsingError(
                "The file contents are not a valid DOCX."
            ) from error

        if "[Content_Types].xml" not in names or "word/document.xml" not in names:
            raise ResumeParsingError("The file contents are not a valid DOCX.")

        if uncompressed_size > settings.max_docx_uncompressed_bytes:
            raise ResumeParsingError("The DOCX contents exceed the allowed limit.")

    @staticmethod
    def _docx_links(document: Document) -> list[ResumeLink]:
        links: list[ResumeLink] = []
        sources = [(document.element.body, document.part)]
        sources.extend(
            (story._element, story.part)
            for story in ResumeParser._docx_header_footers(document)
        )
        for element, part in sources:
            for hyperlink in element.iter(qn("w:hyperlink")):
                relationship_id = hyperlink.get(qn("r:id"))
                if not relationship_id or relationship_id not in part.rels:
                    continue
                relationship = part.rels[relationship_id]
                if not relationship.is_external:
                    continue
                label = "".join(
                    node.text or "" for node in hyperlink.iter(qn("w:t"))
                ).strip()
                links.append(ResumeLink(label=label, url=relationship.target_ref))
        return links

    @staticmethod
    def _docx_header_footers(document: Document):
        seen: set[int] = set()
        for section in document.sections:
            for story in (
                section.header,
                section.first_page_header,
                section.even_page_header,
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            ):
                identifier = id(story.part)
                if identifier not in seen:
                    seen.add(identifier)
                    yield story

    @staticmethod
    def _pdf_link_label(page, link: dict[str, object]) -> str:
        rectangle = link.get("from")
        if not isinstance(rectangle, fitz.Rect):
            return ""
        expanded = rectangle + (-1, -1, 1, 1)
        return re.sub(r"\s+", " ", page.get_textbox(expanded)).strip()

    @classmethod
    def _merge_links(cls, text: str, links: list[ResumeLink]) -> str:
        normalized: list[ResumeLink] = []
        seen: set[str] = set()
        lower_text = text.lower()
        for link in links:
            url = cls._normalize_url(link.url)
            if not url or url.lower() in lower_text or url in seen:
                continue
            seen.add(url)
            normalized.append(ResumeLink(cls._link_label(link.label, url), url))
        if not normalized:
            return text.strip()
        link_lines = [f"{link.label}: {link.url}" for link in normalized]
        return f"{text.strip()}\n\nLINKS\n" + "\n".join(link_lines)

    @staticmethod
    def _normalize_url(value: str) -> str | None:
        value = value.strip()
        if not value:
            return None
        try:
            parsed = urlsplit(value)
        except ValueError:
            return None
        scheme = parsed.scheme.lower()
        if scheme in {"mailto", "tel"}:
            target = parsed.path.strip()
            return f"{scheme}:{target}" if target else None
        if scheme not in {"http", "https"} or not parsed.hostname:
            return None
        hostname = parsed.hostname.lower().rstrip(".")
        if hostname == "localhost" or hostname.endswith(".localhost"):
            return None
        try:
            address = ipaddress.ip_address(hostname)
            if not address.is_global:
                return None
        except ValueError:
            pass
        try:
            port = f":{parsed.port}" if parsed.port else ""
        except ValueError:
            return None
        query = urlencode(
            [
                (key, item)
                for key, item in parse_qsl(parsed.query, keep_blank_values=True)
                if not key.lower().startswith("utm_")
                and key.lower() not in {"fbclid", "gclid"}
            ]
        )
        path = parsed.path.rstrip("/") or ""
        return urlunsplit(("https", f"{hostname}{port}", path, query, ""))

    @staticmethod
    def _link_label(visible_label: str, url: str) -> str:
        if url.startswith("mailto:"):
            return "Email"
        if url.startswith("tel:"):
            return "Phone"
        hostname = (urlsplit(url).hostname or "").lower()
        if hostname == "linkedin.com" or hostname.endswith(".linkedin.com"):
            return "LinkedIn"
        if hostname == "github.com" or hostname.endswith(".github.com"):
            return "GitHub"
        label = re.sub(r"\s+", " ", visible_label).strip()
        if label and len(label) <= 40 and not re.match(r"https?://", label):
            return label
        return "Portfolio"
